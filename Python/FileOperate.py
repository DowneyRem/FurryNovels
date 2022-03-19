#!/usr/bin/python
# -*- coding: UTF-8 -*-
import os
import time
import shutil
from docx.api import Document
from win32com.client import DispatchEx
from functools import wraps


def timethis(func):
	@wraps(func)
	def wrapper(*args, **kwargs):
		start = time.perf_counter()
		r = func(*args, **kwargs)
		end = time.perf_counter()
		print('{}.{} : {}'.format(func.__module__, func.__name__, end - start))
		return r
	return wrapper


pathlist = []
def findFile(path, *extnames):
	for dir in os.listdir(path):
		dir = os.path.join(path, dir)
		if os.path.isdir(dir):
			findFile(dir, *extnames)
		if os.path.isfile(dir):
			for extname in extnames:
				(name, ext) = os.path.splitext(dir)
				if ext == extname:
					pathlist.append(dir)
					
	return pathlist

def openText(path):
	text = ""
	try:
		with open(path, "r", encoding="UTF8") as f:
			text = f.read()
	except UnicodeError:
		try:
			with open(path, "r", encoding="GBK") as f:
				text = f.read()
		except UnicodeError:  # Big5 似乎有奇怪的bug，不过目前似乎遇不到
			with open(path, "r", encoding="BIG5") as f:
				text = f.read()
	finally:
		return text


def openText4(path):
	textlist = []
	try:
		with open(path, "r", encoding="UTF8") as f:
			textlist = f.readlines()[0:4]
	except UnicodeError:
		try:
			with open(path, "r", encoding="GBK") as f:
				textlist = f.readlines()[0:4]
		except UnicodeError:
			with open(path, "r", encoding="BIG5") as f:
				textlist = f.readlines()[0:4]
	finally:
		return textlist

def openDocx(path):
	try:
		docx = Document(path)
		text = ""
		for para in docx.paragraphs:
			if para.style.name == "Normal Indent":  # 正文缩进
				text += "　　" + para.text + "\n"
			else:
				text += para.text + "\n"  # 除正文缩进外的其他所有
		return text
	except IOError:
		print("文件打开失败")
		
		
def openDocx4(path):
	try:
		docx = Document(path)
		text = "";
		j = 1
		for para in docx.paragraphs:
			if j < 5:  # 只读取前4行内容
				j += 1
				if para.style.name == "Normal Indent":  # 正文缩进
					text += "　　" + para.text + "\n"
				else:
					text += para.text + "\n"  # 除正文缩进外的其他所有
			else:
				break
		textlist = text.split("\n")
		return textlist
	except IOError:
		print("文件打开失败")


def makeDirs(path):
	if not os.path.exists(path):
		os.makedirs(path)

def removeFile(path):
	if os.path.exists(path):
		shutil.rmtree(path)
	os.makedirs(path)


def saveDocx(path, text):
	word = DispatchEx('Word.Application')  # 独立进程
	word.Visible = 1  # 0为后台运行
	word.DisplayAlerts = 0  # 不显示，不警告
	template = "D:\\Users\\Administrator\\Documents\\自定义 Office 模板\\小说.dotm"
	docx = word.Documents.Add(template)  # 创建新的word文档
	
	s = word.Selection
	s.Text = text  # 写入文本
	docx.Application.Run("小说排版")  # 运行宏
	
	# 保存文档并退出word
	name = os.path.split(path)[1]
	docx.SaveAs2(path, 16)
	print("【" + name + "】已保存")
	docx.Close(True)
	word.Quit()


def saveText(path, text):
	(dir, name) = os.path.split(path)  # 分离文件名和目录名
	name = name.replace(".txt", "")
	if not os.path.exists(dir):
		os.makedirs(dir)
	try:
		with open(path, "w", encoding="UTF8") as f:
			f.write(text)
		# print("【" + name + "】保存成功")
	except IOError:
		print("【" + name + "】保存失败")


#for循环内部，使用a+模式，写入测试文件
def saveTextDesktop(name, text):
	path = "D:\\Users\\Administrator\\Desktop"
	path = os.path.join(path, name)
	try:
		with open(path, "a+", encoding="UTF8") as f:
			f.write(text)
		# print("【" + name + "】保存成功")
	except IOError:
		print("【" + name + "】保存失败")


def monthNow():
	year = str(time.localtime()[0])
	month = str(time.localtime()[1])
	if len(month) == 1:
		month = "0" + month
	string = os.path.join(year, month)
	return string


def openNowDir():
	path = os.getcwd()
	path = path.replace("\小说推荐\工具", "\兽人小说\小说推荐\频道版")
	text = monthNow()
	path = os.path.join(path,text)
	os.system('start explorer '+ path)


if __name__ == '__main__':
	pass
