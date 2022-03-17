#!/usr/bin/python
# -*- coding: UTF-8 -*-
import os
import time
from docx.api import Document
from opencc import OpenCC
from functools import wraps
from functools import cmp_to_key
cc1 = OpenCC('tw2sp')  #繁转简
cc2 = OpenCC('s2twp')  #簡轉繁


dict = {
		"txt":"txt",
		"docx":"docx",
		"完结":"Finished",
		"完稿":"Finished",
		"finished":"Finished",
		"更新":"Updating",
		"更新中":"Updating",
		"updating":"Updating",
		"停笔":"Died",
		"太监":"Died",
		"died":"Died",
		"简体中文":"zh_cn",
		"简中":"zh_cn",
		"zh_cn":"zh_cn",
		"繁体中文":"zh_tw",
		"正体中文":"zh_tw",
		"繁中":"zh_tw",
		"zh_tw":"zh_tw",
		"中文":"Chinese",
		"汉语":"Chinese",
		"中国语":"Chinese",
		"中国语注意":"Chinese",
		"chinese":"Chinese",
		"原创":"Original",
		"original":"Original",
		"同人":"Doujin",
		"doujin":"Doujin",
		"翻译":"translated",
		"translated":"translated",
  
		"sfw":"SFW",
		"R-18":"R18",
		"R18":"R18",
		"R-18G":"R18G",
		"R18G":"R18G",
		"BL":"Gay",
		"腐向け":"Gay",
		"腐向":"Gay",
		"同志":"Gay",
		"男同性恋":"Gay",
		"gay":"Gay",
		"GL":"Lesbian",
		"百合":"Lesbian",
		"yuri":"Lesbian",
		"女同性恋":"Lesbian",
		"Lesbian":"Lesbian",
		"BG":"General",
		"男女":"General",
		"general":"General",
		"恋爱":"Romance",
		"LOVE":"Romance",
		"romance":"Romance",
		"科幻":"ScienceFiction",
		"ScienceFiction":"ScienceFiction",
		"奇幻":"Fantasy",
		"魔幻":"Fantasy",
		"Fantasy":"Fantasy",
  
		"熊":"bear",
		"bear":"bear",
		"熊猫":"panda",
		"panda":"panda",
		"马":"hourse",
		"hourse":"hourse",
		"牛":"bull",
		"bull":"bull",
		"羊":"sheep",
		"sheep":"sheep",
		"猫":"cat",
		"cat":"cat",
		"狮":"lion",
		"狮子":"lion",
		"lion":"lion",
		"虎":"tiger",
		"老虎":"tiger",
		"tiger":"tiger",
		"西方龙":"dragon",
		"龙":"dragon",
		"龙人":"dragon",
		"dragon":"dragon",
		"东方龙":"long",
		"long":"long",
		"狗":"dog",
		"dog":"dog",
		"狼":"wolf",
		"wolf":"wolf",
		"狐狸":"fox",
		"fox":"fox",
		"鱼":"fish",
		"fish":"fish",
		"鲨鱼":"shark",
		"shark":"shark",
		"海豚":"dolphin",
		"dolphin":"dolphin",
		"大象":"elephant",
		"elephant":"elephant",
		"袋鼠":"kangaroo",
		"kangaroo":"kangaroo",
		"猴子":"monkey",
		"monkey":"monkey",
		"鼠":"mouse",
		"老鼠":"mouse",
		"mouse":"mouse",
		"豹":"panther",
		"panther":"panther",
		"猪":"pig",
		"pig":"pig",
		"兔":"rabbit",
		"兔子":"rabbit",
		"rabbit":"rabbit",
		"犀牛":"rhinoceros",
		"rhinoceros":"rhinoceros",
		"蛇":"snake",
		"snake":"snake",
		"龟":"turtle",
		"turtle":"turtle",
		"鲨狗":"sergal",
		"sergal":"sergal",
  
		"兽人":"Furry",
		"獸人":"Furry",
		"kemono":"Furry",
		"獣人":"Furry",
		"furry":"Furry",
		"纯兽":"non-anthro",
		"non-anthro":"non-anthro",
		"人类":"Human",
		"human":"Human",
		"鸟人":"Harpy",
		"harpy":"Harpy",
		"机器人":"Robot",
		"robot":"Robot",
		"史莱姆":"Slime",
		"slime":"Slime",
		"触手":"Tentacles",
		"tentacles":"Tentacles",
		"吸血鬼":"Vampire",
		"vampire":"Vampire",
		"人皮服装":"SkinSuit",
		"skinsuit":"SkinSuit",
		"怪物":"Monser",
		"monser":"Monser",
		"英雄":"Hero",
		"hero":"Hero",
		"警察":"Police",
		"特警":"Police",
		"police":"Police",
		"魔王":"DevilKing",
		"devilking":"DevilKing",
		"勇者":"Brave",
		"brave":"Brave",
		"魅魔":"Succubus",
		"succubus":"Succubus",
		"医生":"Doctor",
		"doctor":"Doctor",
  
		"阴道交":"VaginalSex",
		"阴道":"VaginalSex",
		"vaginalsex":"VaginalSex",
		"肛交":"Anal",
		"后入":"Anal",
		"anal":"Anal",
		"口交":"BlowJob",
		"blowjob":"BlowJob",
		"生殖腔":"Cloaca",
		"泄殖腔":"Cloaca",
		"cloaca":"Cloaca",
		"尾交":"TailJob",
		"tailjob":"TailJob",
		"自慰":"Masturbation",
		"masturbation":"Masturbation",
		"翼交":"WingJob",
		"wingjob":"WingJob",
		"拳交":"Fisting",
		"fisting":"Fisting",
		"足交":"FootJob",
		"footjob":"FootJob",
		"脚爪":"Paw",
		"足控":"Paw",
		"舔足":"Paw",
		"paw":"Paw",
		"袜交":"SockJob",
		"sockjob":"SockJob",
		"手淫":"HandJob",
		"handjob":"HandJob",
		"脑交":"BrainFuck",
		"brainfuck":"BrainFuck",
		"耳交":"EarFuck",
		"earfuck":"EarFuck",
		"阴茎打脸":"CockSlapping",
		"cockslapping":"CockSlapping",
		"阴茎摩擦":"Frottage",
		"frottage":"Frottage",
		"尿道插入":"UrethraInsertion",
		"urethrainsertion":"UrethraInsertion",
		"奸杀":"RapedWhileDying",
		"rapedwhiledying":"RapedWhileDying",
		"搔痒":"tickling",
		"挠痒":"tickling",
		"tickling":"tickling",
  
		"强奸":"Rape",
		"rape":"Rape",
		"乱伦":"Incest",
		"incest":"Incest",
		"群交":"Group",
		"群p":"Group",
		"轮奸":"Group",
		"group":"Group",
		"绿帽":"Cuckold",
		"netorare":"Cuckold",
		"ntr":"Cuckold",
		"cuckold":"Cuckold",
		"父子丼":"Oyakodon",
		"父子":"Oyakodon",
		"oyakodon":"Oyakodon",
		"兄弟丼":"brother",
		"兄弟":"brother",
		"brother":"brother",
  
		"调教":"BDSM",
		"bdsm":"BDSM",
		"捆绑":"Bondage",
		"束缚":"Bondage",
		"bondage":"Bondage",
		"电击":"ElectricShocks",
		"electricshocks":"ElectricShocks",
		"精神控制":"MindControl",
		"精控":"MindControl",
		"洗脑":"MindControl",
		"mindcontrol":"MindControl",
		"窒息":"Asphyxiation",
		"asphyxiation":"Asphyxiation",
		"高潮禁止":"OrgasmDenial",
		"禁止高潮":"OrgasmDenial",
		"射精控制":"OrgasmDenial",
		"orgasmdenial":"OrgasmDenial",
		"主奴":"Slave",
		"性奴":"Slave",
		"主仆":"Slave",
		"slave":"Slave",
		"鼻吊钩":"NoseHook",
		"nosehook":"NoseHook",
		"狗链":"Leash",
		"leash":"Leash",
		"穿孔":"Piercing",
		"piercing":"Piercing",
		"纹身":"Tattoo",
		"tattoo":"Tattoo",
		"淫纹":"CrotchTattoo",
		"crotchtattoo":"CrotchTattoo",
		"肉便器":"PublicUse",
		"rbq":"PublicUse",
		"publicuse":"PublicUse",
		"饮尿":"PissDrinking",
		"圣水":"PissDrinking",
		"pissdrinking":"PissDrinking",
		"食粪":"Coprophagia",
		"黄金":"Coprophagia",
		"coprophagia":"Coprophagia",
		"人体家具":"Forniphilia",
		"forniphilia":"Forniphilia",
		"人体餐盒":"FoodOnBody",
		"foodonbody":"FoodOnBody",
		"育肥":"WeightGain",
		"肥满化":"WeightGain",
		"weightgain":"WeightGain",
		"战损":"BattleDamage",
		"欠損":"BattleDamage",
		"battledamage":"BattleDamage",
		"虐屌":"CBT",
		"cbt":"CBT",
		"虐腹":"GutTorture",
		"guttorture":"GutTorture",
		"身体改造":"BodyModification",
		"肉体改造":"BodyModification",
		"bodymodification":"BodyModification",
		"变身":"Transformation",
		"transformation":"Transformation",
		"兽化":"Transfur",
		"同化":"Transfur",
		"龙化":"Transfur",
		"transfur":"Transfur",
		"TF":"Transfur",
		"附身":"Possession",
		"possession":"Possession",
		"寄生":"Parasite",
		"parasite":"Parasite",
		"石化":"Petrification",
		"petrification":"Petrification",
		"堕落":"corruption",
		"恶堕":"corruption",
		"corruption":"corruption",
		"催眠":"Hypnosis",
		"hypnosis":"Hypnosis",
		"气味":"Smell",
		"雄臭":"Smell",
		"smell":"Smell",
		"迷药":"Chloroform",
		"chloroform":"Chloroform",
		"药物":"Drugs",
		"药物催淫":"Drugs",
		"drugs":"Drugs",
		"酗酒":"Drunk",
		"drunk":"Drunk",
		"寻欢洞":"GloryHole",
		"gloryhole":"GloryHole",
		"摄像":"Filming",
		"filming":"Filming",
		"性转":"GenderBender",
		"genderbender":"GenderBender",
		"灵魂交换":"BodySwap",
		"身体交换":"BodySwap",
		"bodyswap":"BodySwap",
		"机械奸":"Machine",
		"machine":"Machine",
		"drone":"Drone",
		"尿布":"Diaper",
		"diaper":"Diaper",
  
		"乳胶衣":"Latex",
		"胶衣":"Latex",
		"乳胶":"Latex",
		"latex":"Latex",
		"胶液":"Rubber",
		"胶":"Rubber",
		"龙胶":"Rubber",
		"rubber":"Rubber",
		"生物衣":"LivingClothes",
		"livingclothes":"LivingClothes",
		"防毒面具":"GasMusk",
		"gasmusk":"GasMusk",
		"盔甲":"Armor",
		"armor":"Armor",
		"骑士盔甲":"MetalArmor",
		"metalarmor":"MetalArmor",
		"动力装甲":"PowerArmor",
		"powerarmor":"PowerArmor",
		"军装":"Military",
		"military":"Military",
		"紧身衣":"Leotard",
		"leotard":"Leotard",
		"拘束衣":"StraitJacket",
		"straitjacket":"StraitJacket",
		"泳装":"SwimSuit",
		"swimsuit":"SwimSuit",
		"六尺褌":"Fundoshi",
		"fundoshi":"Fundoshi",
		"袜子":"socks",
		"白袜":"socks",
		"socks":"socks",
  
		"肌肉":"Muscles",
		"筋肉":"Muscles",
		"muscles":"Muscles",
		"巨根":"Hyper",
		"hyper":"Hyper",
		"巨大化":"Macro",
		"macro":"Macro",
		"包茎":"Phimosis",
		"phimosis":"Phimosis",
		"吞食":"Vore",
		"丸吞":"Vore",
		"vore":"Vore",
		"阴茎吞噬":"CockPhagia",
		"cockphagia":"CockPhagia",
		"肛门吞食":"AnalPhagia",
		"analphagia":"AnalPhagia",
		"尾巴吞食":"TailPhagia",
		"tailphagia":"TailPhagia",
		"出产":"Birth",
		"生产":"Birth",
		"birth":"Birth",
		"产卵":"Eggs",
		"eggs":"Eggs",
		"阴茎出产":"PenisBirth",
		"penisbirth":"PenisBirth",
		"肛门出产":"AnalBirth",
		"analbirth":"AnalBirth",
  
		"小说":"",
		}


def timethis(func):
	@wraps(func)
	def wrapper(*args, **kwargs):
		start = time.perf_counter()
		r = func(*args, **kwargs)
		end = time.perf_counter()
		print('{}.{} : {}'.format(func.__module__, func.__name__, end - start))
		return r
	return wrapper


def monthNow():
	year = str(time.localtime()[0])
	month = str(time.localtime()[1])
	if len(month) == 1:
		month = "0" + month
	string = os.path.join(year, month)
	return string


def openNowDir():
	path = os.getcwd()
	path = path.replace("\小说推荐\工具", "\兽人小说\小说推荐\频道版")
	text = monthNow()
	path = os.path.join(path,text)
	os.system('start explorer '+ path)


def cmp(a, b):
	def getindex(a):
		try:
			li = list(dict.values())
			index = li.index(a)
		except:
			li = list(dict.keys())
			index = li.index(a)
		return index
	a = getindex(a)
	b = getindex(b)
	if a > b:
		return 1
	elif a < b:
		return -1
	else:
		return 0


def saveText(path, text):
	(dir, name) = os.path.split(path)  # 分离文件名和目录名
	name = name.replace(".txt", "")
	if not os.path.exists(dir):
		os.makedirs(dir)
	try:
		with open(path, "w", encoding="UTF8") as f:
			f.write(text)
	except IOError:
		print("【" + name + "】保存失败")


def saveDict2md():
	text = "### 关键词标签表\n"
	text = text + "\n| 标签 | 关键词 | "
	text = text + "\n| -- | -- | "
	list1 = list(dict.items())
	for i in range(0, len(list1)):
		(key, value) = list1[i]
		if value != "":
			value = "#" + value
			text += "\n| " + value + " | " + key + " |"
			
	path = os.path.join(os.getcwd(), "Tags.md")
	saveText(path, text)
	
	
def set2text(set):
	text = str(set)
	text = text.replace("{'", "")
	text = text.replace("'}", "")
	text = text.replace("', '", " ")
	text = text.replace(" ", "\n")
	return text


def saveTextDesktop(name, set):
	path = "D:\\Users\\Administrator\\Desktop"
	path = os.path.join(path, name)
	text = set2text(set)
	saveText(path, text)
	
	
def sortTags(set):  # 按dict内顺序对转换后的标签排序
	text = ""
	li = list(set)
	li.sort(key=cmp_to_key(cmp))
	for i in range(len(li)):
		tag = li[i]
		text += "#" + tag + " "
	return text
	
	
def findFile(path):
	for dir in os.listdir(path):
		dir = os.path.join(path, dir)
		if os.path.isdir(dir):
			findFile(dir)
		if os.path.isfile(dir):
			(name, ext) = os.path.splitext(dir)
			if ext == ".docx" or ext == ".txt":
				pathlist.append(dir)
	return pathlist


def openText4(path):
	textlist =[]
	try:
		with open(path, "r", encoding="UTF8") as f:
			textlist = f.readlines()[0:4]
	except UnicodeError:
		try:
			with open(path, "r", encoding="GBK") as f:
				textlist = f.readlines()[0:4]
		except UnicodeError:
			with open(path, "r", encoding = "BIG5") as f:
				textlist = f.readlines()[0:4]
	finally:
		return textlist

	
def openDocx4(path):
	try:
		docx = Document(path)
		text = "";  j = 1
		for para in docx.paragraphs:
			if j < 5:  # 只读取前4行内容
				j += 1
				if para.style.name == "Normal Indent":  # 正文缩进
					text += "　　" + para.text + "\n"
				else:
					text += para.text + "\n"  # 除正文缩进外的其他所有
			else:
				break
		textlist = text.split("\n")
		return textlist
	except IOError:
		print("文件打开失败")


def openText(path):
	text = ""
	try:
		with open(path,"r", encoding = "UTF8") as f:
			text = f.read()
	except UnicodeError:
		try:
			with open(path,"r", encoding = "GBK") as f:
				text = f.read()
		except UnicodeError: #Big5 似乎有奇怪的bug，不过目前似乎遇不到
			with open(path,"r", encoding = "BIG5") as f:
				text = f.read()
	finally:
		return text


def openDocx(path):
	try:
		docx = Document(path)
		text = ""
		for para in docx.paragraphs:
			if para.style.name == "Normal Indent":  # 正文缩进
				text += "　　" + para.text + "\n"
			else:
				text += para.text + "\n"  # 除正文缩进外的其他所有
		return text
	except IOError:
		print("文件打开失败")
		

def addTags(text): #添加靠谱的标签
	list1 = "邊 變 並 從 點 東 對 發 該 個 給 關 過 還 後 歡 會 機 幾 間 見 將 進 經 覺 開 來 裡 兩 嗎 麼 沒 們 難 讓 時 實 說 雖 為 問 無 現 樣 應 於 與 則 這 種".split(" ")
	list2 = "边 变 并 从 点 东 对 发 该 个 给 关 过 还 后 欢 会 机 几 间 见 将 进 经 觉 开 来 里 两 吗 么 没 们 难 让 时 实 说 虽 为 问 无 现 样 应 于 与 则 这 种".split(" ")
	#语料库来自 https://elearning.ling.sinica.edu.tw/cwordfreq.html
	#从中选取前三百的繁体字部分，并在文章中随机检验，取存在率最高的前50个繁体字符
	
	tags = "" ; j = 0 ; list3 = []
	for i in range(len(list1)):
		char = list1[i]
		if char in text:
			j += 1
			list3.append(char)
			
	s1 = set(list1)
	s2 = set(list3)
	s = s1.difference(s2)
	chars = set2text(s)
	k = len(list1) - j
	
	tags += " #txt #finished "
	if j >= 0.2 * len(list1):
		tags += "#zh_tw"
		# print(k)  # 不存在的繁体字符数
		# print(s)   #不存在的繁体字符
	else:
		tags += "#zh_cn"
		if s2 !=set():
			print(s2)  # 存在的繁体字符
	return tags, chars


def translateTags(taglist):  # 获取英文标签
	tags2 = "" ; s = set()
	for i in range(0, len(taglist)):
		tag = taglist[i]
		tag = tag.replace("#", "")
		tag = tag.replace(" ", "")
		tag = tag.replace("　", "")
		tag = dict.get(tag)  #获取英文标签
		
		if tag != None:
			s.add(tag)  # 获取到的标签利用set去重
		else:
			tag = taglist[i]
			tags2 += tag + " "
	tags1 = sortTags(s)  #对转换后的标签排序
	return tags1, tags2
	

def textFormat(textlist, newtags):
	name = cc2.convert(textlist[0]) + "\n"
	authro = textlist[1].replace("作者：", "")
	authro = "by #" + authro + "\n"
	
	url = textlist[2].replace("网址：", "")
	url = url.replace("網址：", "")
	url = url.replace("链接：", "")
	url = url + "\n"
	
	tags = textlist[3].replace("标签：", "")
	tags = tags.replace("標簽：", "")
	tags += newtags  #新增 #zh_tw 或 #zh_cn
	tags = cc1.convert(tags)  #转简体，只处理简体标签
	list = tags.split()
	(tags1, tags2) = translateTags(list)  #获取已翻译/未翻译的标签
	
	text = name + authro + tags1 + "\n特殊：" + tags2 + "\n" + url #+ "\n"
	print(text)   #格式化输出
	return tags2  #输出不支持的标签


def printTags(path):
	pathlist = findFile(path)
	dirstr = monthNow()  # 只处理本月的文件
	s = set() ; j = 0 ; chars =""
	for i in range(0, len(pathlist)):
		path = pathlist[i]
		(dir, name) = os.path.split(path)
		(name, ext) = os.path.splitext(name)
		if dirstr in dir:  # 只处理本月的文件
			j += 1
			if ext == ".docx":
				textlist = openDocx4(path)
				text     = openDocx (path)
			elif ext == ".txt":
				textlist = openText4(path)
				text     = openText (path)
				
			if j >= 0:  #无用的if语句，保持上下几行缩进一致
				newtags, char = addTags(text) #根据本文繁简添加标签
				chars += char +"\n"*1
				s.add(textFormat(textlist, newtags))
	
	saveTextDesktop("tags.txt", s)
	# saveTextDesktop("文字.txt", chars)
	
	if j != 0:
		# openNowDir()
		pass
	else:
		print("本月 " + dirstr + " 无新文档")
	
	
def main():
	print("本月文档如下：")
	print("\n"*2)
	printTags(path)
	saveDict2md()
	
	
if __name__ == "__main__":
	path = os.path.join(os.getcwd())
	path = path.replace("\工具", "")
	pathlist = []
	main()
