#!/usr/bin/python
# -*- coding: UTF-8 -*-
import os
import re
import json
import time
import shutil
import logging
import webbrowser
from functools import wraps
from platform import platform

# import zipfile as zf
import pyzipper as zf
import pdfplumber
from docx import Document  # 使用 docx-hitalent

from configuration import testMode


logging.basicConfig(
		level=logging.INFO,
		format='%(levelname)s %(asctime)s [%(filename)s:%(lineno)d] %(message)s',
		datefmt='%Y.%m.%d. %H:%M:%S',
		# filename='parser_result.log',
		# filemode='w'
		)


template_docx = os.path.join("data", "模板.docx")
template_dotx = os.path.join("data", "模板.dotx")
template_dotm = os.path.join("data", "模板.dotm")
office_old_ext = ".doc .ppt".split(" ")  # xls 可由 xlrd==1.2.0 读取


def timer(function):
	@wraps(function)
	def wrapper(*args, **kwargs):
		start = time.perf_counter()
		result = function(*args, **kwargs)
		end = time.perf_counter()
		print(f"{function.__module__}.{function.__name__}: {end - start}")
		return result
	return wrapper


def onWindows(function):
	@wraps(function)
	def wrapper(*args, **kwargs):
		result = ""
		if "Windows" in platform():  # 导入 pywin32 特有库
			try:
				global Dispatch, DispatchEx, winreg
				import winreg
				from win32com.client import Dispatch, DispatchEx
			except (ModuleNotFoundError, ImportError) as e:
				print("不存在: pywin32")
				logging.error(e)
			else:
				result = function(*args, **kwargs)
		else:
			logging.error(f"非 Windows 不可执行 {function}")
		return result
	return wrapper
	
	
def openFileCheck(function):
	@wraps(function)
	def wrapper(*args, **kwargs):
		path = args[0]
		if os.path.exists(path):
			try:
				result = function(*args, **kwargs)
				return result
			except IOError:
				logging.error(f"文件被占用：{path}")
		else:
			logging.error(f"文件不存在：{path}")
	return wrapper


def saveFileCheck(function):
	@wraps(function)
	def wrapper(*args, **kwargs):
		os.makedirs(os.path.dirname(args[0]), exist_ok=True)
		result = function(*args, **kwargs)
		return result
	return wrapper


@onWindows
def desktop() -> str:
	if "Windows" in platform():  # 其他平台没用过
		key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
		path = winreg.QueryValueEx(key, "Desktop")[0]
	else:  # 未测试
		path = os.path.expanduser("~/Desktop")
	return path


def makeDirs(path: str, delete=False):
	if not os.path.exists(path):
		os.makedirs(path)
	if os.path.exists(path) and delete:  # 删除后重建
		removeFile(path)
		os.makedirs(path)
		# os.makedirs(path, exist_ok=True)


def makeFile(path: str, data=""):  # 新建空文件，不检查拓展名
	if not os.path.exists(path):
		return saveText(path, data)
	
	
def copyFile(path1: str, path2: str):
	shutil.copy2(path1, path2)
	
	
def removeFile(*paths: str):
	# 删除多个文件或文件夹
	for path in paths:
		if os.path.isdir(path):
			try:
				shutil.rmtree(path)
				logging.info(f"已经删除：{path}")
			except IOError:
				logging.error(f"删除失败：{path}")
		
		elif os.path.isfile(path):
			try:
				os.remove(path)
				logging.info(f"已经删除：{path}")
			except IOError:
				logging.error(f"删除失败：{path}")


def formatFileName(text: str) -> str:
	if text:
		text = re.sub('[:*?"<>|]', ' ', text)
		text = text.replace(os.sep, '')
		text = text.replace('  ', '')
		return text
	else:
		return ""


def findFile(path: str, *extnames: str) -> list:
	"""
	Args:
		path: path 需要遍历的文件夹路径
		extnames: extname=""获取无后缀名文件；省略 extnames 参数可以获取全部文件
	"""
	pathlist = []
	for root, dirs, files in os.walk(path):
		# print(root, dirs, files, sep="\n")
		for file in files:
			fullpath = os.path.join(root, file)
			
			if len(extnames) > 0:
				for extname in extnames:
					if fullpath.lower().endswith(extname):
						pathlist.append(fullpath)
			elif len(extnames) == 0:
				pathlist.append(fullpath)
	return pathlist


def openFile(path: str):
	# 使用默认应用打开文件
	if os.path.exists(path):
		webbrowser.open(path)
	
	
@onWindows
@openFileCheck
def openExcel(path: str):  # 打开 Excel
	app = DispatchEx('Excel.Application')  # 独立进程
	app.Visible = 1  # 0为后台运行
	app.DisplayAlerts = 0  # 不显示，不警告
	try:
		app.Workbooks.Open(path)  # 打开文档
		print("打开Excel……")
	except IOError as e:
		logging.error(e)


def readFile(path: str):
	# 根据文件后缀名，自动选用读取函数
	result = ""
	extname = os.path.splitext(path)[1]
	funname = f"read{extname.replace(os.path.extsep, '').capitalize()}"
	if not extname:
		result = readText(path)
	else:
		try:
			result = globals()[funname](path)
		except KeyError:
			logging.error(f"没有 {funname} 方法")
			print(f"没有 {funname} 方法")
	# print(result)
	return result


@openFileCheck
def readText(path: str) -> str:
	text = ""
	try:
		with open(path, "r", encoding="UTF8") as f:
			text = f.read()
	except UnicodeError:
		try:
			with open(path, "r", encoding="GB18030") as f:
				text = f.read()
		except UnicodeError:  # Big5 似乎有奇怪的bug，不过目前似乎遇不到
			try:
				with open(path, "r", encoding="BIG5") as f:
					text = f.read()
			except UnicodeError:
				with open(path, "r", encoding="Latin1") as f:
					text = f.read()
	finally:
		return text


readTxt = readText
readMd = readText


@openFileCheck
def readDocx(path: str) -> str:
	text = ""
	try:
		docx = Document(path)
	except IOError as e:
		logging.error(e)
	else:
		for para in docx.paragraphs:
			# print(para.paragraph_format)
			# print(para.paragraph_format.first_line_indent.pt)
			if para.style.name == "Normal Indent":  # 正文缩进
				text += f"　　{para.text}\n"
			elif para.paragraph_format.first_line_indent and para.paragraph_format.first_line_indent.pt >= 15:  # 首行缩进
				text += f"　　{para.text}\n"
			else:
				text += f"{para.text}\n"
	return text


@timer
@onWindows
@openFileCheck
def readDoc(path: str) -> str:
	text = ""
	word = DispatchEx('Word.Application')  # 独立进程
	word.Visible = 0  # 0为后台运行
	word.DisplayAlerts = 0  # 不显示，不警告
	
	try:
		docx = word.Documents.Open(path)
	except IOError as e:
		logging.error(e)
	else:
		text = docx.Content.Text
		text = text.replace("\r\r", "\n").replace("\x0b", "\n")  # 换行符；手动换行符
		if docx.Paragraphs.CharacterUnitFirstLineIndent == 2:  # 首行缩进转空格
			textlist = ["", ] + text.split()
			text = "\n　　".join(textlist).strip("\n")
		# print(len(text), text, sep="\n")
		docx.Close(True)
	finally:
		word.Quit()
		return text


@openFileCheck
def readJson(path: str) -> any:
	try:
		with open(path, "rb") as f:
			data = json.load(f)
	except IOError as e:
		logging.error(e)
	else:
		return data
	
	
def readENPdf(path: str, *, password="", **kwargs) :
	name = os.path.basename(path)
	with pdfplumber.open(path, password=password) as pdf:
		newpages = []
		for page in pdf.pages:
			text = page.extract_text()
			lines = text.split("\n")  # 获取每一行的文本
			# print(text)
			newlines = []
			for line in lines:
				# print(line)
				words = line.split(" ")
				# print(words)
				
				# 段落首行
				if len(words) > 1 and words[0][0].isupper():  # 首字母大写
					# print(words)
					if len(words) > 1 and words == lines[0].split(" "):  # 第一行
						count = 0
						l1 = " ".join(words).lower().split()
						l2 = name.lower().replace(".", " ").replace("_", " ").split()
						for word in l1:
							if word in l2:
								count += 1
								
						if count >= len(l1) - 1: # 是标题
						# if " ".join(words).lower() in name.split("_"):
							words = words + ["\n"]
						else:
							words = ["\n", "   "] + words    # 3个空格
							
					else:
						words = ["\n", "   "] + words    # 3个空格
					# print(words)
					
				# 段落末行
				if len(words) > 1 and words[-1][-1] in ". ? ! 。 ？ ！".split():
					words.extend(["\n", "   "])  # 3个空格
				# print(words)
				newline = " ".join(words)  # 空格间隔每个词
				# print(newline)
				newlines.append(newline)
			# print(newlines)
			newpage = "".join(newlines)
			# print(newpage)
			newpages.append(newpage)
		# print(newpage)
		newtext = "".join(newpages)
		# print(newtext)
	return newtext


def readCNPdf(path: str, *, password="", **kwargs):
	name = os.path.basename(path)
	with pdfplumber.open(path, password=password) as pdf:
		# first_page = pdf.pages[0]
		# print(first_page.extract_text())
		newpages = []
		for page in pdf.pages:
			text = page.extract_text()
			lines = text.split("\n")  # 获取每一行的文本
			# print(text)
			newlines = []
			for line in lines:
				words = line.split()  # 拆分每行内容
				# print(words)
				if len(words) > 0 and words == lines[0].split():  # 第一行
					# print(words)
					if name in words or words[0] in name:  # 第一行是标题
						words = words + ["\n\n", "　　"]
					else:
						words = ["\n", "　　"] + words
					
				if len(words) > 0 and words[-1][-1] in ". ? ! 。 ？ ！".split():  # 最后一行最后一个字符
					words.extend(["\n", "　　"])
				# print(words)
				newline = "".join(words)
				# print(newline)
				newlines.append(newline)
			# print(newlines)
			newpage = "".join(newlines)
			# print(newpage)
			newpages.append(newpage)
		# print(newpage)
		newtext = "".join(newpages)
		# print(newtext)
	return newtext


@openFileCheck
def readPdf(path: str, *, password="", **kwargs) -> any:
	with pdfplumber.open(path, password=password) as pdf:
		first_page = pdf.pages[0].extract_text()
		# first_page = pdf.pages[0].extract_words()
		# print(first_page)
		
	if "，" in first_page or "。" in first_page:
		text = readCNPdf(path, password=password, **kwargs)
	else:
		text = readENPdf(path, password=password, **kwargs)
	if not text:
		raise ValueError(f"不支持扫描版 PDF 文件")
	return text


@saveFileCheck
def saveFile(path: str, data: any, **kwargs):
	# 根据文件后缀名，自动选用保存函数
	extname = os.path.splitext(path)[1].lower()
	if not extname:
		return saveText(path, data)
	elif extname in office_old_ext:
		path, extname = f"{path}x", f"{extname}x"  # .doc -> .docx
		print(f"保存为新格式：{path}")
	elif extname == ".pdf":
		extname = ".txt"
		path = f"{os.path.splitext(path)[0]}{extname}"
		
	funname = f"save{extname.replace(os.path.extsep, '').capitalize()}"
	try:
		globals()[funname](path, data, **kwargs)
	except KeyError:  # 有后缀名，无函数
		logging.error(f"没有 {funname} 方法")
	return path


@saveFileCheck
def saveText(path: str, text="", **kwargs):
	try:
		with open(path, "w", encoding="UTF8") as f:
			f.write(text)
	except IOError:
		logging.error(f"保存失败：{path}")
	else:
		logging.debug(f"已保存为：{path}")
	return path


saveMd = saveText
saveCsv = saveText
saveIni = saveText
saveTxt = saveText


@saveFileCheck
def saveTextDesktop(name: str, text: str):
	path = os.path.join(desktop(), name)
	return saveText(path, text)


@saveFileCheck
def saveDocx(path: str, text: str, *, template="", **kwargs):
	"""
	Args:
		path: path 待压缩的文件/文件夹路径
		text: text 文本
		template: path 模板路径，docx 格式
	"""
	if os.path.exists(template) and template.lower().endswith(".docx"):  # 使用模板
		docx = Document(template)
		for para in docx.paragraphs:  # 清空
			para.clear()
			para._element.getparent().remove(para._element)
	else:    # 使用 data 目录下指定模板
		docx = Document(template_docx)
	
	for para in text.split("\n"):
		docx.add_paragraph(para)
	try:
		docx.save(path)
	except IOError as e:
		logging.error(e)
	else:
		logging.debug(f"已保存为：{path}")
	return path


@onWindows
@saveFileCheck
def saveDoc(path: str, text: str, **kwargs):
	extdict = {
		".docx": 16,
		# ".docm": "",  # 启用宏的文档
		".dotx": 1,  # Word 模板
		# ".dotm": "",  # 启用宏的模板
		".doc": 0,  # 97-03 .doc
		".pdf": 17,
		".xps": 18,
		".txt": 7,  # Unicode text
		".odt": 23, # OpenDocument Text format
	}
	
	app = DispatchEx('Word.Application')  # 独立进程
	app.Visible = 0  # 0为后台运行
	app.DisplayAlerts = 0  # 不显示，不警告
	try:
		doc = app.Documents.Add(template_dotm)  # 创建新的word文档
	except IOError as e:
		logging.error(e)
	else:
		# text = docx.Content.Text
		s = app.Selection
		s.FormatText = text  # 写入文本
		doc.Application.Run("小说排版")  # 运行宏
		extname = os.path.splitext(path)[1]
		if extname == ".txt":
			doc.SaveAs2(path, 7, Encoding=65001, AddToRecentFiles=0, AllowSubstitutions=0, LineEnding=0)  # txt UTF8 CRLF
		else:
			doc.SaveAs2(path, extdict.get(extname, 16))
		logging.debug(f"已保存为：{path}")
		doc.Close(True)
	finally:
		app.Quit()
	return path

	
@saveFileCheck
def saveJson(path: str, data: any, **kwargs):
	try:
		with open(path, 'w', encoding="UTF8") as f:
			json.dump(data, f, ensure_ascii=False, indent=4)
	except IOError:
		logging.error(f"保存失败：{path}")
	else:
		logging.debug(f"已保存为：{path}")
	return path

	
def zipMultiFiles(path: [str, list], *, password="", zippath="", delete=0) -> str:
	"""
	压缩多个的文件成单个zip文件；使用 pyzipper 可用aes256加密
	Args:
		path: path 待压缩的文件路径
		password: password 密码
		zippath: path 指定zip路径
		delete: delete != 0 时，删除源文件
	"""
	if isinstance(path, str):  # 单个文件
		path = [path]
		folder = os.path.dirname(path[0])  # 上级文件夹
	else:
		folder = os.path.commonpath(path)  # 共同路径，应该是上级文件夹
	
	if password:
		encryption = zf.WZ_AES
	else:
		encryption = None
	
	if zippath:
		os.makedirs(os.path.dirname(zippath), exist_ok=True)
	else:
		if len(path) == 1:
			zippath = f"{os.path.splitext(path[0])[0]}.zip"
		else:
			zippath = f"{os.path.split(path[0])[0]}.zip"
	
	removeFile(zippath)
	# with zf.ZipFile(zippath, 'w', compression=zf.ZIP_DEFLATED) as z:
	with zf.AESZipFile(zippath, 'w', compression=zf.ZIP_LZMA, encryption=encryption) as z:
		z.setpassword(password.encode(encoding="utf-8"))
		# z.comment = b"Comment"  # 无法写入注释
		
		for file in path:
			arcname = os.path.relpath(file, folder)  # 获取 zip 内文件路径
			# print(filepath, arcname, sep="\n")
			z.write(filename=file, arcname=arcname)
	print(f"压缩完成：{zippath}")
	
	if delete:
		for file in path:
			removeFile(file)
	return zippath
	
	
def zipEachFile(path: [str, list], *, password="", zippath="", delete=0) -> str:
	"""
	添加到单独”压缩文件名.zip“，压缩多个文件成多个zip文件；使用 pyzipper 可用aes256加密
	Args:
		path: path 待压缩的文件路径
		password: password 密码
		zippath: path 指定zip路径，或其存放其的文件夹
		delete: delete != 0 时，删除源文件
	"""
	if isinstance(path, str):  # 单个文件
		path = [path]
		
	if password:
		encryption = zf.WZ_AES
	else:
		encryption = None
	
	_zippath = ""  # 区分传进来的 zippath；真正的压缩后的路径
	if zippath:
		if zippath.lower().endswith(".zip"):  # 将指定的zip文件路径转换成文件夹路径
			zipfolder = os.path.dirname(zippath)
		else:
			zipfolder = zippath
		os.makedirs(zipfolder, exist_ok=True)
	else:
		zipfolder = ""
		
	for file in path:
		folder = os.path.dirname(file)  # 上级文件夹
		if zippath and zipfolder:  # 指定路径写入压缩文件
			if len(path) == 1:  # 单个文件，丢失中间路径
				name = os.path.basename(file)
				_zippath = os.path.join(zipfolder, f"{os.path.splitext(name)[0]}.zip")
			else:  # 多个文件，保留中间路径
				name = os.path.relpath(file, os.path.commonpath(path))
				_zippath = os.path.join(zipfolder, f"{os.path.splitext(name)[0]}.zip")
		else:  # 源路径写入压缩文件
			_zippath = f"{os.path.splitext(file)[0]}.zip"
		os.makedirs(os.path.dirname(_zippath), exist_ok=True)
		removeFile(_zippath)
		
		# with zf.ZipFile(_zippath, 'w', compression=zf.ZIP_DEFLATED) as z:
		with zf.AESZipFile(_zippath, 'w', compression=zf.ZIP_LZMA, encryption=encryption) as z:
			z.setpassword(password.encode(encoding="utf-8"))
			# z.comment = b"Comment"  # 无法写入注释
			arcname = os.path.relpath(file, folder)  # 获取 zip 内文件路径
			# print(filepath, arcname, sep="\n")
			z.write(filename=file, arcname=arcname)
		# print(f"压缩完成：{arcname}")
		print(f"压缩完成：{_zippath}")
		
		if delete:
			removeFile(file)
	return _zippath


def zipFolder(path: str, *, password="", zippath="", delete=0) -> str:
	"""
	压缩文件夹与子文件；使用 pyzipper 可用aes256加密
	Args:
		path: path 待压缩的文件夹路径
		password: password 密码
		zippath: path 指定zip路径
		delete: delete != 0 时，删除源文件
	"""
	if password:
		encryption = zf.WZ_AES
	else:
		encryption = None
	
	if not zippath:
		zippath = f"{path}.zip"
	removeFile(zippath)
	# with zf.ZipFile(zippath, 'w', compression=zf.ZIP_DEFLATED) as z:
	with zf.AESZipFile(zippath, 'w', compression=zf.ZIP_LZMA, encryption=encryption) as z:
		z.setpassword(password.encode(encoding="utf-8"))
		
		files = findFile(path, )  # 获取目录下所有文件
		folder = os.path.dirname(path)  # 上级文件夹
		for file in files:
			arcname = os.path.relpath(file, folder)  # 替换上级文件夹，建立当前文件夹
			# print(filepath, arcname, sep="\n")
			z.write(filename=file, arcname=arcname)
	print(f"压缩完成：{zippath}")
	
	if delete:
		removeFile(path)
	return zippath


def zipFile(path: [str, list], *, password="", zippath="", delete=0, each=0) -> str:
	"""
	压缩传入的文件或文件夹；使用 pyzipper 可用aes256加密
	Args:
		path: path 待压缩的文件/文件夹路径
		password: password 密码
		zippath: path 指定zip路径
		delete: delete != 0 时，删除源文件
		each: each==1 时，将相应文件添加到单独”压缩文件名.zip“
	"""
	if isinstance(path, list) or os.path.isfile(path):
		if each:
			zippath = zipEachFile(path, password=password, zippath=zippath, delete=delete)
		else:
			zippath = zipMultiFiles(path, password=password, zippath=zippath, delete=delete)
	elif os.path.isdir(path):
		zippath = zipFolder(path, password=password, zippath=zippath, delete=delete)
	else:
		print(f"无法压缩，不存在：{path}")
		return ""
	return zippath


def unzipZipFileEasy(zippath: str, *, password="", mode=1, delete=0) -> str:
	"""
	解压 zip 文件；使用 pyzipper 可解压加密的zip文件（ase256 与 ZipCrypto）,前者会快得多
	智能解压：zip内无文件夹，或有多个文件夹，新建以zip文件名为名的文件夹，再解压
	智能解压：zip内只有1个文件夹，且与zip文件相同，直接解压
	常规软件压缩设置：勾选zip使用Unicode文件名，避免解压后文件名乱码
	Args:
		zippath: path 待解压的zip文件
		password: password 密码
		mode: mode=1 解压zip内部的zip文件；同时删除内部zip
		delete: delete=1 解压后删除zip源文件
	"""
	
	# with zf.ZipFile(zippath, "r") as z:
	with zf.AESZipFile(zippath, "r") as z:
		if z.namelist()[0].endswith("/") or len(z.namelist()) == 1:  # 单文件，内有文件夹，直接解压
			folder = os.path.split(zippath)[0]
		else:  # 多个文件，新建文件夹
			folder = os.path.splitext(zippath)[0]
		
		try:
			for file in z.namelist():
				z.extract(file, folder, password.encode('utf-8'))
				# if not file.endswith("/"):
				# 	print(f"解压：{file}")
				
				if file.endswith(".zip") and mode:  # 解压zip内的zip
					_zippath = os.path.join(folder, file)
					unzipZipFileEasy(_zippath, password=password, mode=mode, delete=1)
			# print(f"解压：{file}")
			print(f"解压完成：{zippath}")
		except RuntimeError:
			print(f"密码【{password}】错误，解压失败：{zippath}")
	
	if delete != 0:
		removeFile(zippath)  # 删除zip文件
	return folder


def unzipZipFile(zippath: str, *, password="", path="", mode=1, delete=0) -> str:
	"""
	解压 zip 文件；使用 pyzipper 可解压加密的zip文件（ase256 与 ZipCrypto）,前者会快得多
	智能解压：zip内无文件夹，或有多个文件夹，新建以zip文件名为名的文件夹，再解压
	智能解压：zip内只有1个文件夹，且与zip文件相同，直接解压
	常规软件压缩设置：勾选zip使用Unicode文件名，避免解压后文件名乱码
	Args:
		zippath: path 待解压的zip文件
		password: password 密码
		path: path 指定路径解压
		mode: mode=1 解压zip内部的zip文件；同时删除内部zip
		delete: delete=1 解压后删除zip源文件
	"""
	# with zf.ZipFile(zippath, "r") as z:
	with zf.AESZipFile(zippath, "r") as z:
		if z.comment:
			comment = z.comment.decode(encoding="utf-8")
			print(f"压缩文件注释：{comment}")
		
		zip_name = os.path.splitext(os.path.basename(zippath))[0].strip()
		if len(z.namelist()) >= 2:
			zip_dir0 = os.path.commonpath(z.namelist()).strip()
		else:
			# zip_dir0 = z.namelist()[0].replace("/", "")
			zip_dir0 = os.path.basename(z.namelist()[0]).strip()
		
		if path:  # 指定位置解压
			os.makedirs(path, exist_ok=True)
			if zip_dir0 == zip_name:  # zip 内文件夹与zip名称相同
				result = os.path.join(path, zip_dir0)  # 返回路径
			else:  # zip 内有多个文件，新建文件夹
				path = result = os.path.join(path, zip_name)  # 解压到的目录 & 返回路径
		else:  # 默认位置解压
			result = os.path.splitext(zippath)[0].strip()  # 返回路径
			if zip_dir0 == zip_name:  # zip 内文件夹与zip名称相同
				path = os.path.dirname(zippath).strip()  # 解压到的目录
			else:  # zip 内有多个文件，新建文件夹
				path = os.path.splitext(zippath)[0].strip()  # 解压到的目录
		
		try:
			for file in z.namelist():  # 解压zip
				z.extract(file, path, password.encode('utf-8'))
				if file.lower().endswith(".zip") and mode:  # 解压zip内的zip
					_zippath = os.path.join(path, file)
					print(f"解压：{_zippath}")
					unzipZipFile(_zippath, password=password, mode=mode, delete=1)
			print(f"解压完成：{zippath}")
			# print(f"文件目录：{result}")
		except RuntimeError:
			print(f"密码【{password}】错误，解压失败：{zippath}")
	
	if delete != 0:
		removeFile(zippath)  # 删除zip文件
	return result


def unzipFolder(zippath: str, *, password="", path="", mode=0, delete=0) -> str:
	"""
	使用 pyzipper 可解压加密的zip文件（ase256 与 ZipCrypto）,前者会快得多
	常规软件压缩设置：勾选zip使用Unicode文件名，避免解压后文件名乱码
	Args:
		zippath: path 含有zip文件的文件夹路径
		password: password 密码
		path: path 指定路径解压
		mode: mode=1 解压zip内部的zip文件
		delete: delete=1 解压后删除zip源文件；同时 mode=1 解压后会删除所有zip
	"""
	zippaths = findFile(zippath, ".zip")
	if len(zippaths) == 0:
		print(f"{zippath} 目录下无zip文件")
		return ""
	
	file = ""
	for zipfile in zippaths:  # 未测试 path 参数
		print(f"正在解压：{zipfile}")
		file = unzipZipFile(zipfile, password=password, path=path, mode=mode, delete=delete)
		# print("————" * 30, "", sep="\n")
	return file


def unzipFile(zippath: str, *, password="", path="", mode=0, delete=0) -> str:
	"""
	使用 pyzipper 可解压加密的zip文件（ase256 与 ZipCrypto）,前者会快得多
	智能解压：path传入zip路径解压zip，传入文件夹则解压其路径下的zip
	智能解压：zip内无文件夹，则会新建以zip文件名为名的文件夹，zip只有单文件不新建文件夹
	常规软件压缩设置：勾选zip使用Unicode文件名，避免解压后文件名乱码
	Args:
		zippath: path 待解压的zip文件/含有zip文件的文件夹路径
		password: password 密码
		path: path 指定路径解压
		mode: mode=1 解压zip内部的zip文件
		delete: delete=1 解压后删除zip源文件；同时 mode=1 解压后会删除所有zip
	"""
	if zf.is_zipfile(zippath):
		return unzipZipFile(zippath, password=password, path=path, mode=mode, delete=delete)
	elif os.path.isdir(zippath):
		return unzipFolder(zippath, password=password, mode=mode, delete=delete)
	else:
		print(f"{path}非zip文件，亦非含有zip文件的目录")
		return ""


def test():
	print("测试")
	

if __name__ == '__main__':
	test()
