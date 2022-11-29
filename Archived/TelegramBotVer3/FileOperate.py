#!/usr/bin/python
# -*- coding: UTF-8 -*-
import os
import re
import json
import time
import shutil
import logging
import pathlib
import webbrowser
from functools import wraps
from platform import platform

# import zipfile as zf
import pyzipper as zf
from docx import Document  # 使用 docx-hitalent

from config import testMode


logging.basicConfig(
		level=logging.INFO,
		format='%(levelname)s %(asctime)s [%(filename)s:%(lineno)d] %(message)s',
		datefmt='%Y.%m.%d. %H:%M:%S',
		# filename='parser_result.log',
		# filemode='w'
		)


template_docx = os.path.join("data", "模板.docx")
template_dotx = os.path.join("data", "模板.dotx")
template_dotm = os.path.join("data", "模板.dotm")
office_old_ext = ".doc .ppt".split(" ")  # xls 可由 xlrd==1.2.0 读取


def timer(function):
	@wraps(function)
	def wrapper(*args, **kwargs):
		start = time.perf_counter()
		result = function(*args, **kwargs)
		end = time.perf_counter()
		print(f"{function.__module__}.{function.__name__}: {end - start}")
		return result
	return wrapper


def checkWindows(function):
	@wraps(function)
	def wrapper(*args, **kwargs):
		result = ""
		if "Windows" in platform():  # 导入 pywin32 特有库
			try:
				global Dispatch, DispatchEx, winreg
				import winreg
				from win32com.client import Dispatch, DispatchEx
			except (ModuleNotFoundError, ImportError) as e:
				print("不存在: pywin32")
				logging.error(e)
			else:
				result = function(*args, **kwargs)
		else:
			logging.error(f"非 Windows 不可执行 {function}")
		return result
	return wrapper
	
	
def openFileCheck(function):
	@wraps(function)
	def wrapper(*args, **kwargs):
		path = args[0]
		if os.path.exists(path):
			try:
				result = function(*args, **kwargs)
				return result
			except IOError:
				logging.error(f"文件被占用：{path}")
		else:
			logging.error(f"文件不存在：{path}")
	return wrapper


def saveFileCheck(function):
	@wraps(function)
	def wrapper(*args, **kwargs):
		makeDirs(os.path.dirname(args[0]))
		result = function(*args, **kwargs)
		return result
	return wrapper


@checkWindows
def desktop() -> str:
	if "Windows" in platform():  # 其他平台没用过
		key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
		path = winreg.QueryValueEx(key, "Desktop")[0]
	else:  # 未测试
		path = os.path.expanduser("~/Desktop")
	return path


def makeDirs(path: str):
	if not os.path.exists(path):
		os.makedirs(path)
	# else:
	# 	removeFile(path)
	# 	makeDirs(path)
	
	
def makeFile(path: str, data=""):  # 新建空文件，不检查拓展名
	if not os.path.exists(path):
		return saveText(path, data)
	
	
def copyFile(path1: str, path2: str):
	shutil.copy2(path1, path2)
	
	
def removeFile(*paths: str):
	# 删除多个文件或文件夹
	for path in paths:
		if os.path.isdir(path):
			try:
				shutil.rmtree(path)
				logging.info(f"已经删除：{path}")
			except IOError:
				logging.error(f"删除失败：{path}")
		
		elif os.path.isfile(path):
			try:
				os.remove(path)
				logging.info(f"已经删除：{path}")
			except IOError:
				logging.error(f"删除失败：{path}")


def formatFileName(text: str) -> str:
	if text:
		text = re.sub('[:*?"<>|]', ' ', text)
		text = text.replace(os.sep, '')
		text = text.replace('  ', '')
		return text
	else:
		return ""


def findFile(path: str, *extnames: str) -> list:
	"""
	Args:
		path: path 需要遍历的文件夹路径
		extnames: extname=""获取无后缀名文件；省略 extnames 参数可以获取全部文件
	"""
	pathlist = []
	for root, dirs, files in os.walk(path):
		# print(root, dirs, files, sep="\n")
		for file in files:
			fullpath = os.path.join(root, file)
			
			if len(extnames) > 0:
				for extname in extnames:
					if fullpath.lower().endswith(extname):
						pathlist.append(fullpath)
			elif len(extnames) == 0:
				pathlist.append(fullpath)
	return pathlist


def openFile(path: str):
	# 使用默认应用打开文件
	if os.path.exists(path):
		webbrowser.open(path)
	
	
@checkWindows
@openFileCheck
def openExcel(path: str):  # 打开 Excel
	app = DispatchEx('Excel.Application')  # 独立进程
	app.Visible = 1  # 0为后台运行
	app.DisplayAlerts = 0  # 不显示，不警告
	try:
		app.Workbooks.Open(path)  # 打开文档
		print("打开Excel……")
	except IOError as e:
		logging.error(e)


def readFile(path: str):
	# 根据文件后缀名，自动选用读取函数
	result = ""
	extname = os.path.splitext(path)[1]
	funname = f"read{extname.replace(os.path.extsep, '').capitalize()}"
	if not extname:
		result = readText(path)
	else:
		try:
			result = globals()[funname](path)
		except KeyError:
			logging.error(f"没有 {funname} 方法")
			print(f"没有 {funname} 方法")
	# print(result)
	return result


@openFileCheck
def readText(path: str) -> str:
	text = ""
	try:
		with open(path, "r", encoding="UTF8") as f:
			text = f.read()
	except UnicodeError:
		try:
			with open(path, "r", encoding="GB18030") as f:
				text = f.read()
		except UnicodeError:  # Big5 似乎有奇怪的bug，不过目前似乎遇不到
			try:
				with open(path, "r", encoding="BIG5") as f:
					text = f.read()
			except UnicodeError:
				with open(path, "r", encoding="Latin1") as f:
					text = f.read()
	finally:
		return text


readTxt = readText
readMd = readText


@openFileCheck
def readDocx(path: str) -> str:
	text = ""
	try:
		docx = Document(path)
	except IOError as e:
		logging.error(e)
	else:
		for para in docx.paragraphs:
			# print(para.paragraph_format)
			# print(para.paragraph_format.first_line_indent.pt)
			if para.style.name == "Normal Indent":  # 正文缩进
				text += f"　　{para.text}\n"
			elif para.paragraph_format.first_line_indent and para.paragraph_format.first_line_indent.pt >= 15:  # 首行缩进
				text += f"　　{para.text}\n"
			else:
				text += f"{para.text}\n"
	return text


@timer
@checkWindows
@openFileCheck
def readDoc(path: str) -> str:
	text = ""
	word = DispatchEx('Word.Application')  # 独立进程
	word.Visible = 0  # 0为后台运行
	word.DisplayAlerts = 0  # 不显示，不警告
	
	try:
		docx = word.Documents.Open(path)
	except IOError as e:
		logging.error(e)
	else:
		text = docx.Content.Text
		text = text.replace("\r\r", "\n").replace("\x0b", "\n")  # 换行符；手动换行符
		if docx.Paragraphs.CharacterUnitFirstLineIndent == 2:  # 首行缩进转空格
			textlist = ["", ] + text.split()
			text = "\n　　".join(textlist).strip("\n")
		# print(len(text), text, sep="\n")
		docx.Close(True)
	finally:
		word.Quit()
		return text


@openFileCheck
def readJson(path: str) -> any:
	try:
		with open(path, "rb") as f:
			data = json.load(f)
	except IOError as e:
		logging.error(e)
	else:
		return data
	

@saveFileCheck
def saveFile(path: str, data: any, **kwargs):
	# 根据文件后缀名，自动选用保存函数
	extname = os.path.splitext(path)[1]
	if not extname:
		return saveText(path, data)
	elif extname in office_old_ext:
		path, extname = f"{path}x", f"{extname}x"  # .doc -> .docx
		print(f"保存为新格式：{path}")
	funname = f"save{extname.replace(os.path.extsep, '').capitalize()}"
	try:
		globals()[funname](path, data, **kwargs)
	except KeyError:  # 有后缀名，无函数
		logging.error(f"没有 {funname} 方法")


@saveFileCheck
def saveText(path: str, text="", **kwargs):
	try:
		with open(path, "w", encoding="UTF8") as f:
			f.write(text)
	except IOError:
		logging.error(f"保存失败：{path}")
	else:
		logging.debug(f"已保存为：{path}")


saveMd = saveText
saveCsv = saveText
saveIni = saveText
saveTxt = saveText


@saveFileCheck
def saveTextDesktop(name: str, text: str):
	path = os.path.join(desktop(), name)
	return saveText(path, text)


@saveFileCheck
def saveDocx(path: str, text: str, *, template="", **kwargs):
	"""
	Args:
		path: path 待压缩的文件/文件夹路径
		text: text 文本
		template: path 模板路径，docx 格式
	"""
	if os.path.exists(template) and template.lower().endswith(".docx"):  # 使用模板
		docx = Document(template)
		for para in docx.paragraphs:  # 清空
			para.clear()
			para._element.getparent().remove(para._element)
	else:    # 使用 data 目录下指定模板
		docx = Document(template_docx)
	
	for para in text.split("\n"):
		docx.add_paragraph(para)
	try:
		docx.save(path)
	except IOError as e:
		logging.error(e)
	else:
		logging.debug(f"已保存为：{path}")


@checkWindows
@saveFileCheck
def saveDoc(path: str, text: str, **kwargs):
	extdict = {
		".docx": 16,
		# ".docm": "",  # 启用宏的文档
		".dotx": 1,  # Word 模板
		# ".dotm": "",  # 启用宏的模板
		".doc": 0,  # 97-03 .doc
		".pdf": 17,
		".xps": 18,
		".txt": 7,  # Unicode text
		".odt": 23, # OpenDocument Text format
	}
	
	app = DispatchEx('Word.Application')  # 独立进程
	app.Visible = 0  # 0为后台运行
	app.DisplayAlerts = 0  # 不显示，不警告
	try:
		doc = app.Documents.Add(template_dotm)  # 创建新的word文档
	except IOError as e:
		logging.error(e)
	else:
		# text = docx.Content.Text
		s = app.Selection
		s.FormatText = text  # 写入文本
		doc.Application.Run("小说排版")  # 运行宏
		extname = os.path.splitext(path)[1]
		if extname == ".txt":
			doc.SaveAs2(path, 7, Encoding=65001, AddToRecentFiles=0, AllowSubstitutions=0, LineEnding=0)  # txt UTF8 CRLF
		else:
			doc.SaveAs2(path, extdict.get(extname, 16))
		logging.debug(f"已保存为：{path}")
		doc.Close(True)
	finally:
		app.Quit()
	
	
@saveFileCheck
def saveJson(path: str, data: any, **kwargs):
	try:
		with open(path, 'w', encoding="UTF8") as f:
			json.dump(data, f, ensure_ascii=False, indent=4)
	except IOError:
		logging.error(f"保存失败：{path}")
	else:
		logging.debug(f"已保存为：{path}")
	
	
def zipMultiFiles(path: [str, list], *, password="", zippath="", delete=0) -> str:
	"""压缩单个或多个的文件；使用 pyzipper 可用aes256加密
	Args:
		path: path 待压缩的文件路径
		password: password 密码
		zippath: path 指定zip路径
		delete: delete != 0 时，删除源文件
	"""
	if password:
		encryption = zf.WZ_AES
	else:
		encryption = None
		
	if isinstance(path, str):  # 单个文件
		path = [path]
		folder = os.path.dirname(path[0])  # 上级文件夹
		if not zippath:
			zippath = f"{os.path.splitext(path[0])[0]}.zip"
	else:
		folder = os.path.commonpath(path)  # 共同路径，应该是上级文件夹
		if not zippath:
			zippath = f"{os.path.split(path[0])[0]}.zip"
	
	removeFile(zippath)
	# with zf.ZipFile(zippath, 'w', compression=zf.ZIP_DEFLATED) as z:
	with zf.AESZipFile(zippath, 'w', compression=zf.ZIP_LZMA, encryption=encryption) as z:
		z.setpassword(password.encode(encoding="utf-8"))
		# z.comment = b"Comment"  # 无法写入注释
		
		for file in path:
			arcname = os.path.relpath(file, folder)  # 获取 zip 内文件路径
			# print(filepath, arcname, sep="\n")
			z.write(filename=file, arcname=arcname)
	
	if delete:
		for file in path:
			removeFile(file)
	return zippath
	
	
def zipFolder(path: str, *, password="", zippath="", delete=0) -> str:
	"""压缩文件夹与子文件；使用 pyzipper 可用aes256加密
	Args:
		path: path 待压缩的文件夹路径
		password: password 密码
		zippath: path 指定zip路径
		delete: delete != 0 时，删除源文件
	"""
	if password:
		encryption = zf.WZ_AES
	else:
		encryption = None
	
	if not zippath:
		zippath = f"{path}.zip"
	removeFile(zippath)
	# with zf.ZipFile(zippath, 'w', compression=zf.ZIP_DEFLATED) as z:
	with zf.AESZipFile(zippath, 'w', compression=zf.ZIP_LZMA, encryption=encryption) as z:
		z.setpassword(password.encode(encoding="utf-8"))
		
		files = findFile(path, )  # 获取目录下所有文件
		folder = os.path.dirname(path)  # 上级文件夹
		for file in files:
			arcname = os.path.relpath(file, folder)  # 替换上级文件夹，建立当前文件夹
			# print(filepath, arcname, sep="\n")
			z.write(filename=file, arcname=arcname)
	
	if delete:
		removeFile(path)
	return zippath


def zipFile(path: [str, list], *, password="", zippath="", delete=0) -> str:
	"""压缩传入的文件或文件夹；使用 pyzipper 可用aes256加密
	Args:
		path: path 待压缩的文件/文件夹路径
		password: password 密码
		zippath: path 指定zip路径
		delete: delete != 0 时，删除源文件
	"""
	if isinstance(path, list) or os.path.isfile(path):
		zippath = zipMultiFiles(path, password=password, zippath=zippath, delete=delete)
	elif os.path.isdir(path):
		zippath = zipFolder(path, password=password, zippath=zippath, delete=delete)
	else:
		print(f"无法压缩，不存在：{path}")
		return ""
	print(f"压缩完成：{zippath}")
	return zippath
	
	
def unzipZipFile(zippath: str, *, password="", path="", mode=1, delete=0) -> str:
	"""解压 zip 文件；使用 pyzipper 可解压加密的zip文件（ase256 与 ZipCrypto）,前者会快得多
	智能解压：zip内无文件夹，或有多个文件夹，新建以zip文件名为名的文件夹，再解压
	智能解压：zip内只有1个文件夹，且与zip文件相同，直接解压
	常规软件压缩设置：勾选zip使用Unicode文件名，避免解压后文件名乱码
	Args:
		zippath: path 待解压的zip文件
		password: password 密码
		path: path 指定路径解压
		mode: mode=1 解压zip内部的zip文件
		delete: delete=1 解压后删除zip源文件；同时 mode=1 解压后会删除所有zip
	"""
	# with zf.ZipFile(path, "r") as z:
	with zf.AESZipFile(zippath, "r") as z:
		if z.comment:
			comment = z.comment.decode(encoding="utf-8")
			print(f"压缩文件注释：{comment}")
		
		zip_name = os.path.splitext(os.path.basename(zippath))[0]
		if len(z.namelist()) >= 2:
			zip_dir0 = os.path.commonpath(z.namelist())
		else:
			zip_dir0 = z.namelist()[0].replace("/", "")
		
		if path:  # 指定位置解压
			makeDirs(path)
			# zip 内文件夹与zip名称相同
			if zip_dir0 == zip_name:
				result = os.path.join(path, zip_dir0)  # 返回路径
			else:  # zip 内有多个文件，新建文件夹
				path = result = os.path.join(path, zip_name)  # 解压到的目录 & 返回路径
				
		else:  # 默认位置解压
			result = os.path.splitext(zippath)[0]  # 返回路径
			# zip 内文件夹与zip名称相同
			if zip_dir0 == zip_name:
				path = os.path.dirname(zippath)  # 解压到的目录
			else:  # zip 内有多个文件，新建文件夹
				path = os.path.splitext(zippath)[0]  # 解压到的目录
		
		try:
			for file in z.namelist():  # 解压zip
				z.extract(file, path, password.encode('utf-8'))
				if file.lower().endswith(".zip") and mode:  # 解压zip内的zip
					zippath = os.path.join(path, file)
					print("解压：", zippath)
					unzipZipFile(zippath, password=password, mode=mode, delete=delete)
			print(f"解压完成：{zippath}")
			# print(f"文件目录：{result}")
		except RuntimeError:
			print(f"密码【{password}】错误，解压失败：{zippath}")
	
	if delete != 0:
		removeFile(zippath)  # 删除zip文件
	return result


def unzipFolder(zippath: str, *, password="", path="", mode=0, delete=0) -> str:
	"""使用 pyzipper 可解压加密的zip文件（ase256 与 ZipCrypto）,前者会快得多
	常规软件压缩设置：勾选zip使用Unicode文件名，避免解压后文件名乱码
	Args:
		zippath: path 含有zip文件的文件夹路径
		password: password 密码
		path: path 指定路径解压
		mode: mode=1 解压zip内部的zip文件
		delete: delete=1 解压后删除zip源文件；同时 mode=1 解压后会删除所有zip
	"""
	ziplist = findFile(zippath, ".zip")
	if len(ziplist) == 0:
		print(f"{zippath} 目录下无zip文件")
		return ""
	
	for zipfile in ziplist:  # 未测试 path 参数
		return unzipZipFile(zipfile, password=password, path=path, mode=mode, delete=delete)


def unzipFile(zippath: str, *, password="", path="", mode=0, delete=0) -> str:
	"""使用 pyzipper 可解压加密的zip文件（ase256 与 ZipCrypto）,前者会快得多
	智能解压：path传入zip路径解压zip，传入文件夹则解压其路径下的zip
	智能解压：zip内无文件夹，则会新建以zip文件名为名的文件夹，zip只有单文件不新建文件夹
	常规软件压缩设置：勾选zip使用Unicode文件名，避免解压后文件名乱码
	Args:
		zippath: path 待解压的zip文件/含有zip文件的文件夹路径
		password: password 密码
		path: path 指定路径解压
		mode: mode=1 解压zip内部的zip文件
		delete: delete=1 解压后删除zip源文件；同时 mode=1 解压后会删除所有zip
	"""
	if zf.is_zipfile(zippath):
		return unzipZipFile(zippath, password=password, path=path, mode=mode, delete=delete)
	elif os.path.isdir(zippath):
		return unzipFolder(zippath, password=password, mode=mode, delete=delete)
	else:
		print(f"{path}非zip文件，亦非含有zip文件的目录")
		return ""


def test():
	print("测试")
	

if __name__ == '__main__':
	test()
